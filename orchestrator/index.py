# orchestrator/index.py
import os, sys, re, json, hashlib, argparse, pathlib, time
from typing import List, Dict, Any

ROOT = pathlib.Path(__file__).resolve().parents[1]  # repo root
CFG  = ROOT / "orchestrator" / "docs_config.json"
IDX_DIR = ROOT / "ops" / "index"
MANIFEST = IDX_DIR / "manifest.jsonl"
SEGMENTS = IDX_DIR / "segments.jsonl"
TFIDF_PKL = IDX_DIR / "tfidf.pkl"
REPORT = ROOT / "docs" / "reports" / "index-coverage.json"

def sha256_bytes(data: bytes) -> str:
    import hashlib; h=hashlib.sha256(); h.update(data); return h.hexdigest()

def load_config() -> Dict[str, Any]:
    with open(CFG, "r", encoding="utf-8") as f:
        return json.load(f)

def should_include(path: pathlib.Path, inc: List[str], exc: List[str]) -> bool:
    from fnmatch import fnmatch
    sp = str(path).replace("\\","/")
    ok = any(fnmatch(sp, (ROOT / gl).as_posix()) or fnmatch(sp, gl) for gl in inc) if inc else True
    if not ok: return False
    if exc and any(fnmatch(sp, (ROOT / gx).as_posix()) or fnmatch(sp, gx) for gx in exc): return False
    return True

def read_text(path: pathlib.Path) -> str:
    sp = str(path).lower()
    try:
        if sp.endswith((".md",".txt",".py",".yml",".yaml",".json")):
            return path.read_text("utf-8", errors="ignore")
        if sp.endswith(".docx"):
            try:
                import docx
                doc = docx.Document(str(path))
                return "\n".join(p.text for p in doc.paragraphs)
            except Exception:
                return f"[[UNPARSED: {path.name} (docx)]]"
        if sp.endswith(".odt"):
            try:
                from odfdo import Document
                d = Document(str(path))
                return d.get_formatted_text("\n")
            except Exception:
                return f"[[UNPARSED: {path.name} (odt)]]"
        return path.read_text("utf-8", errors="ignore")
    except Exception as e:
        return f"[[READ_ERROR: {path.name}: {e}]]"

def chunk_text(text: str, size=1200, overlap=200):
    if not text: return []
    tokens = text.splitlines()
    out, i, n = [], 0, len(tokens)
    while i < n:
        j = min(n, i + size)
        out.append("\n".join(tokens[i:j]))
        i = j - overlap if (j - overlap) > i else j
    return out

def scan_and_build(_):
    cfg = load_config()
    IDX_DIR.mkdir(parents=True, exist_ok=True)
    (ROOT / "docs" / "reports").mkdir(parents=True, exist_ok=True)

    for f in (MANIFEST, SEGMENTS):
        if f.exists(): f.unlink()

    max_bytes = int(cfg.get("max_file_size_mb", 8)) * 1024 * 1024
    csize, coverlap = cfg["chunk"]["size"], cfg["chunk"]["overlap"]

    files = []
    for root in cfg["roots"]:
        base = ROOT / root["path"]
        include = root.get("include", [])
        exclude = root.get("exclude", [])
        for p in base.rglob("*"):
            if not p.is_file(): continue
            if p.stat().st_size > max_bytes: continue
            if should_include(p, include, exclude):
                files.append((root["name"], p))

    with MANIFEST.open("w", encoding="utf-8") as mf, SEGMENTS.open("w", encoding="utf-8") as sf:
        for group, path in files:
            b = path.read_bytes()
            meta = {
                "group": group,
                "path": str(path.relative_to(ROOT)),
                "bytes": len(b),
                "sha256": sha256_bytes(b),
                "mtime": int(path.stat().st_mtime)
            }
            mf.write(json.dumps(meta) + "\n")

            text = read_text(path)
            for i, part in enumerate(chunk_text(text, size=csize, overlap=coverlap), start=0):
                # track start/end by rough line math within chunks
                lines = part.count("\n") + 1
                start = 1 + i * (csize - coverlap if csize > coverlap else csize)
                seg = {
                    "group": group,
                    "path": meta["path"],
                    "start": start,
                    "end": start + lines - 1,
                    "text": part[:20000]
                }
                sf.write(json.dumps(seg) + "\n")

    # Optional TF‑IDF
    try:
        from sklearn.feature_extraction.text import TfidfVectorizer
        import pickle
        texts = []
        with SEGMENTS.open("r", encoding="utf-8") as f:
            for line in f: texts.append(json.loads(line)["text"])
        vec = TfidfVectorizer(max_features=50000, ngram_range=(1,2))
        X = vec.fit_transform(texts)
        with open(TFIDF_PKL, "wb") as pf:
            pickle.dump({"vec": vec, "X": X}, pf)
    except Exception:
        pass

def iter_segments():
    with SEGMENTS.open("r", encoding="utf-8") as f:
        for line in f: yield json.loads(line)

def query(args):
    q, k = args.text, args.k
    results = []
    try:
        import pickle, numpy as np
        with open(TFIDF_PKL, "rb") as pf:
            obj = pickle.load(pf)
        vec, X = obj["vec"], obj["X"]
        qv = vec.transform([q])
        scores = (X @ qv.T).toarray().ravel()
        top = scores.argsort()[::-1][:k]
        segs = list(iter_segments())
        for i in top:
            if i < len(segs):
                s = segs[i].copy(); s["score"] = float(scores[i]); results.append(s)
    except Exception:
        # naive fallback
        for s in iter_segments():
            if re.search(re.escape(q), s["text"], re.IGNORECASE):
                s["score"] = 1.0; results.append(s)
        results = results[:k]

    for s in results:
        print(f"{s['path']}:{s['start']}-{s['end']}  score={s.get('score',0):.4f}")
        snip = s["text"].strip().splitlines()
        print("> " + ("\n> ".join(snip[:6]))[:800]); print()

def read_cmd(args):
    p = (ROOT / args.path).resolve()
    if not p.exists():
        print(f"NOT FOUND: {args.path}", file=sys.stderr); sys.exit(2)
    text = read_text(p)
    lines = text.splitlines()
    a = 1 if args.first else args.start
    b = (args.first or args.end or len(lines))
    a = max(a,1); b = min(b, len(lines))
    print("```text"); print("\n".join(lines[a-1:b])); print("```")

def audit(args):
    # Use curated blueprint indexes if present
    required = []
    for f in ["docs/blueprints/nt-index.json", "docs/blueprints/td-index.json"]:
        path = ROOT / f
        if path.exists():
            try:
                data = json.loads(path.read_text("utf-8"))
                if isinstance(data, list):
                    for row in data:
                        if isinstance(row, str): required.append(row)
                        elif isinstance(row, dict) and "path" in row: required.append(row["path"])
            except Exception:
                pass

    present = set()
    if MANIFEST.exists():
        with MANIFEST.open("r", encoding="utf-8") as mf:
            for line in mf: present.add(json.loads(line)["path"])

    missing = [r for r in required if r not in present]
    report = {
        "ts": int(time.time()),
        "manifest_exists": MANIFEST.exists(),
        "segments_exists": SEGMENTS.exists(),
        "required_total": len(required),
        "required_missing": missing,
        "present_total": len(present)
    }
    REPORT.parent.mkdir(parents=True, exist_ok=True)
    REPORT.write_text(json.dumps(report, indent=2), "utf-8")

    if args.strict and (not MANIFEST.exists() or not SEGMENTS.exists() or missing):
        print("AUDIT FAIL", json.dumps(report, indent=2)); sys.exit(1)
    else:
        print("AUDIT OK", json.dumps(report, indent=2))

def main():
    ap = argparse.ArgumentParser("orchestrator.index")
    sp = ap.add_subparsers(dest="cmd", required=True)
    b = sp.add_parser("build"); b.set_defaults(func=scan_and_build)
    q = sp.add_parser("query"); q.add_argument("--text", required=True); q.add_argument("--k", type=int, default=8); q.set_defaults(func=query)
    r = sp.add_parser("read");  r.add_argument("--path", required=True); r.add_argument("--first", type=int, default=0); r.add_argument("--start", type=int, default=1); r.add_argument("--end", type=int, default=0); r.set_defaults(func=read_cmd)
    a = sp.add_parser("audit"); a.add_argument("--strict", action="store_true"); a.set_defaults(func=audit)
    args = ap.parse_args(); args.func(args)

if __name__ == "__main__":
    main()
