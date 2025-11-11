from pathlib import Path
from docx import Document
from odf.opendocument import load as odf_load
from odf import text as odf_text, table as odf_table

NT = Path("docs/blueprints/Combined_Master_PLAIN_Non_Tech_001.docx")
TD = Path("docs/blueprints/TechnicalDevelopmentPlan.odt")

def docx_blocks(p: Path, limit=40):
    if not p.exists():
        print(f"(missing) {p}"); return
    doc = Document(str(p))
    out = []
    # paragraphs
    for i, para in enumerate(doc.paragraphs):
        t = para.text.strip()
        if t:
            style = getattr(para.style, "name", "")
            out.append(("P", i, style, t))
        if len(out) >= limit: break
    # tables (first few cells if still under limit)
    if len(out) < limit:
        ti = 0
        for tbl in doc.tables:
            for r in tbl.rows:
                for c in r.cells:
                    t = c.text.strip()
                    if t:
                        out.append((f"T{ti}", None, "cell", t))
                        ti += 1
                        if len(out) >= limit: break
                if len(out) >= limit: break
            if len(out) >= limit: break
    print(f"\nDOCX peek: {p}  (showing {len(out)} non-empty blocks)")
    for tag, idx, style, t in out:
        head = t if len(t) <= 120 else t[:117] + "..."
        print(f"[{tag}:{idx if idx is not None else '-'} | {style}] {head}")

def odt_blocks(p: Path, limit=40):
    if not p.exists():
        print(f"(missing) {p}"); return
    odt = odf_load(str(p))
    out = []
    # headings
    for i, h in enumerate(odt.getElementsByType(odf_text.H)):
        txt = "".join(n.data for n in h.childNodes if n.nodeType == n.TEXT_NODE).strip()
        if txt:
            out.append(("H", i, "", txt))
        if len(out) >= limit: break
    # paragraphs (if room remains)
    if len(out) < limit:
        for i, ptag in enumerate(odt.getElementsByType(odf_text.P)):
            txt = "".join(n.data for n in ptag.childNodes if n.nodeType == n.TEXT_NODE).strip()
            if txt:
                out.append(("P", i, "", txt))
            if len(out) >= limit: break
    # list items (if room remains)
    if len(out) < limit:
        for li in odt.getElementsByType(odf_text.ListItem):
            ps = li.getElementsByType(odf_text.P)
            for ptag in ps:
                txt = "".join(n.data for n in ptag.childNodes if n.nodeType == n.TEXT_NODE).strip()
                if txt:
                    out.append(("L", None, "", txt))
                if len(out) >= limit: break
            if len(out) >= limit: break
    print(f"\nODT peek: {p}  (showing {len(out)} non-empty blocks)")
    for tag, idx, _, t in out:
        head = t if len(t) <= 120 else t[:117] + "..."
        print(f"[{tag}:{idx if idx is not None else '-'}] {head}")

if __name__ == "__main__":
    docx_blocks(NT, limit=40)
    odt_blocks(TD, limit=40)
